"""
DAMA IA Service - GSM Buscador de Editais v42.0

Motor de Inteligência para geração automática de propostas comerciais.

Pipeline:
1. OCR: Extrai dados do PDF do edital usando Gemini 2.5 Flash
2. Análise: Identifica itens, valores, órgão, modalidade
3. Cálculo Tributário: Aplica regras de ICMS interestadual
4. Geração: Cria proposta usando template Word do cliente
5. Saída: ZIP com Proposta (.docx) + Declarações

Integração: emergentintegrations com Gemini 2.5 Flash
"""

import os
import io
import zipfile
import tempfile
import logging
from datetime import datetime, timezone
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass
from docx import Document
from docx.shared import Pt, Inches
import re

from dotenv import load_dotenv
load_dotenv()

logger = logging.getLogger(__name__)


# ==================== CONFIGURAÇÃO TRIBUTÁRIA ICMS ====================

ICMS_INTERESTADUAL = {
    # Origem -> Destino -> Alíquota
    # Sul/Sudeste (12, 7) para outras regiões
    'SP': {'SP': 18, 'RJ': 12, 'MG': 12, 'ES': 12, 'PR': 12, 'SC': 12, 'RS': 12, 'DEFAULT': 7},
    'RJ': {'SP': 12, 'RJ': 20, 'MG': 12, 'ES': 12, 'PR': 12, 'SC': 12, 'RS': 12, 'DEFAULT': 7},
    'MG': {'SP': 12, 'RJ': 12, 'MG': 18, 'ES': 12, 'PR': 12, 'SC': 12, 'RS': 12, 'DEFAULT': 7},
    'PR': {'SP': 12, 'RJ': 12, 'MG': 12, 'ES': 12, 'PR': 18, 'SC': 12, 'RS': 12, 'DEFAULT': 7},
    'SC': {'SP': 12, 'RJ': 12, 'MG': 12, 'ES': 12, 'PR': 12, 'SC': 17, 'RS': 12, 'DEFAULT': 7},
    'RS': {'SP': 12, 'RJ': 12, 'MG': 12, 'ES': 12, 'PR': 12, 'SC': 12, 'RS': 18, 'DEFAULT': 7},
    'ES': {'SP': 12, 'RJ': 12, 'MG': 12, 'ES': 17, 'PR': 12, 'SC': 12, 'RS': 12, 'DEFAULT': 7},
    # Norte/Nordeste/Centro-Oeste (12% para Sul/Sudeste, interno varia)
    'DEFAULT': {'SP': 12, 'RJ': 12, 'MG': 12, 'ES': 12, 'PR': 12, 'SC': 12, 'RS': 12, 'DEFAULT': 12}
}


@dataclass
class EmpresaGSM:
    """Dados de uma empresa do Grupo Smart Medical"""
    id: str
    nome: str
    cnpj: str
    endereco: str
    uf_origem: str
    inscricao_estadual: str
    telefone: str
    email: str


# Empresas cadastradas do Grupo Smart Medical v52.0
# c1 = HC IMPORTAÇÕES (usa modelo HC.docx)
# Demais = VIP FARMA ou outras (usa modelo Vip.docx)
EMPRESAS_GSM = {
    'c1': EmpresaGSM(
        id='c1',
        nome='HC IMPORTAÇÕES EXPORTAÇÕES LTDA',
        cnpj='31.958.700/0001-17',
        endereco='Rua Domingos Dadalto, 127, Galpão 03, Rio Branco, Cariacica - ES',
        uf_origem='ES',
        inscricao_estadual='084.050.99-3',
        telefone='(11) 3164-4607',
        email='claudio@gruposmartmedical.com.br'
    ),
    'c2': EmpresaGSM(
        id='c2',
        nome='VIP FARMA COMÉRCIO DE MEDICAMENTOS EIRELI',
        cnpj='34.788.645/0001-52',
        endereco='Rua Arapaçu, 90, Sala 04, Vila Formosa, São Paulo/SP - CEP 03358-000',
        uf_origem='SP',
        inscricao_estadual='146.887.780.119',
        telefone='(11) 3164-4607',
        email='licitacao@vipfarma.com.br'
    ),
    'c3': EmpresaGSM(
        id='c3',
        nome='Smart Medical - Filial SC',
        cnpj='22.222.222/0001-22',
        endereco='Av. Central, 500 - Florianópolis/SC - CEP 88000-000',
        uf_origem='SC',
        inscricao_estadual='987.654.321.000',
        telefone='(48) 3000-0000',
        email='filialsc@smartmedical.com.br'
    ),
    'c4': EmpresaGSM(
        id='c4',
        nome='GSM Distribuidora Eireli',
        cnpj='33.333.333/0001-33',
        endereco='Rua Industrial, 200 - Campinas/SP - CEP 13000-000',
        uf_origem='SP',
        inscricao_estadual='111.222.333.444',
        telefone='(19) 3000-0000',
        email='distribuidora@gsm.com.br'
    )
}

# Mapeamento de template por empresa v52.0
# c1 (HC IMPORTAÇÕES) -> modelo HC.docx
# Demais -> modelo Vip.docx
TEMPLATE_MAP = {
    'c1': 'HC',       # HC IMPORTAÇÕES -> modelo HC.docx
    'default': 'VIP'  # Demais empresas -> modelo Vip.docx
}


@dataclass
class DadosEdital:
    """Dados extraídos do edital via OCR"""
    orgao: str
    cnpj_orgao: str
    endereco_orgao: str
    uf_destino: str
    numero_processo: str
    modalidade: str
    objeto: str
    itens: List[Dict]
    data_abertura: str
    valor_referencia_total: float
    observacoes: str


class DAMAService:
    """
    Serviço principal do módulo DAMA IA.
    
    Responsável por:
    - OCR de PDFs de editais
    - Extração de dados estruturados
    - Cálculo tributário
    - Geração de propostas em Word
    - Montagem do Kit de Licitação (ZIP)
    """
    
    def __init__(self):
        self.api_key = os.environ.get('EMERGENT_LLM_KEY')
        if not self.api_key:
            logger.warning("⚠️ EMERGENT_LLM_KEY não configurada")
    
    async def processar_edital(
        self,
        pdf_content: bytes,
        docx_template_content: bytes,
        empresa_id: str,
        custo_unitario: float = 0.0,
        empresa_data: Optional[Dict] = None
    ) -> Tuple[bytes, Dict]:
        """
        Pipeline completo de processamento do edital (v46.0).
        
        Args:
            pdf_content: Conteúdo binário do PDF do edital
            docx_template_content: Conteúdo binário do template Word
            empresa_id: ID da empresa do Grupo GSM
            custo_unitario: Custo unitário informado pelo usuário
            empresa_data: Dados da empresa vindos do banco (opcional)
            
        Returns:
            Tuple[bytes, Dict]: (ZIP com documentos, estatísticas)
        """
        try:
            logger.info("🚀 [DAMA] Iniciando processamento do edital...")
            
            # Validar empresa - priorizar dados do banco
            if empresa_data and empresa_data.get("nome"):
                # Criar empresa a partir dos dados do banco
                empresa = EmpresaGSM(
                    id=empresa_id,
                    nome=empresa_data.get("nome", ""),
                    cnpj=empresa_data.get("cnpj", ""),
                    endereco=empresa_data.get("endereco", ""),
                    uf_origem=self._extrair_uf_endereco(empresa_data.get("endereco", "")),
                    inscricao_estadual=empresa_data.get("ie", ""),
                    telefone=empresa_data.get("telefone", ""),
                    email=empresa_data.get("email", "")
                )
                logger.info(f"🏢 [DAMA] Usando empresa do banco: {empresa.nome}")
            else:
                empresa = EMPRESAS_GSM.get(empresa_id)
                if not empresa:
                    raise ValueError(f"Empresa não encontrada: {empresa_id}")
            
            # 1. OCR e extração de dados
            logger.info("📄 [DAMA] Etapa 1: OCR do PDF...")
            dados_edital = await self._extrair_dados_pdf(pdf_content)
            
            # 2. Gerar texto da proposta (v73.1: sem cálculo tributário)
            logger.info("✍️ [DAMA] Etapa 2: Gerando proposta...")
            texto_proposta = await self._gerar_texto_proposta_simples(
                dados_edital,
                empresa
            )
            
            # 3. Injetar no template Word
            logger.info("📝 [DAMA] Etapa 3: Injetando no template Word...")
            proposta_docx = self._injetar_template(
                docx_template_content,
                texto_proposta,
                dados_edital,
                empresa
            )
            
            # 4. Gerar declarações obrigatórias
            logger.info("📋 [DAMA] Etapa 4: Gerando declarações...")
            declaracoes = self._gerar_declaracoes(empresa, dados_edital, docx_template_content)
            
            # 5. Montar ZIP
            logger.info("📦 [DAMA] Etapa 5: Montando ZIP...")
            zip_content = self._montar_zip(proposta_docx, declaracoes, dados_edital)
            
            stats = {
                'orgao': dados_edital.orgao,
                'numero_processo': dados_edital.numero_processo,
                'itens_processados': len(dados_edital.itens),
                'valor_referencia': dados_edital.valor_referencia_total,
                'empresa': empresa.nome
            }
            
            logger.info(f"✅ [DAMA] Processamento concluído! {stats}")
            
            return zip_content, stats
            
        except Exception as e:
            logger.error(f"❌ [DAMA] Erro no processamento: {str(e)}")
            raise
    
    def _extrair_uf_endereco(self, endereco: str) -> str:
        """
        Extrai a UF de um endereço (v46.0).
        
        Ex: "Rua X, 123, Cariacica - ES" -> "ES"
        """
        # Lista de UFs válidas
        ufs = ['AC','AL','AP','AM','BA','CE','DF','ES','GO','MA','MT','MS','MG','PA','PB','PR','PE','PI','RJ','RN','RS','RO','RR','SC','SP','SE','TO']
        
        # Tentar encontrar UF no formato "- UF" ou ", UF"
        for uf in ufs:
            if f"- {uf}" in endereco.upper() or f", {uf}" in endereco.upper() or f" {uf}." in endereco.upper():
                return uf
            if endereco.upper().endswith(f" {uf}") or endereco.upper().endswith(f"-{uf}"):
                return uf
        
        # Default
        return 'SP'
    
    async def extrair_itens_edital(self, pdf_content: bytes) -> List[Dict]:
        """
        v61.0: Extrai apenas os itens do edital para configuração prévia.
        
        Retorna lista de itens com campos para o usuário configurar:
        - item: número do item
        - descricao: descrição técnica
        - quantidade: quantidade licitada
        - unidade: unidade de medida
        - valor_referencia: valor de referência (se houver)
        """
        from emergentintegrations.llm.chat import LlmChat, UserMessage, FileContentWithMimeType
        
        # Salvar PDF temporariamente
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
            tmp.write(pdf_content)
            pdf_path = tmp.name
        
        try:
            chat = LlmChat(
                api_key=self.api_key,
                session_id=f"dama-itens-{datetime.now().timestamp()}",
                system_message="""Você é um especialista em análise de editais de licitação.
Sua tarefa é extrair APENAS OS ITENS do edital.

SEMPRE responda em formato JSON válido com a seguinte estrutura:
{
    "itens": [
        {
            "item": "1",
            "descricao": "Descrição técnica completa do item",
            "quantidade": 100,
            "unidade": "UN",
            "valor_referencia": 10.50
        }
    ]
}

REGRAS:
- Extraia TODOS os itens do edital
- Se o valor de referência não existir, use 0
- Unidades comuns: UN, CX, FR, PCT, KG, L, ML, M, M2, M3
- Mantenha a descrição técnica completa"""
            ).with_model("gemini", "gemini-2.5-flash")
            
            pdf_file = FileContentWithMimeType(
                file_path=pdf_path,
                mime_type="application/pdf"
            )
            
            response = await chat.send_message(UserMessage(
                text="Extraia TODOS os itens deste edital de licitação em formato JSON.",
                file_contents=[pdf_file]
            ))
            
            import json
            
            response_text = response.strip()
            if response_text.startswith('```'):
                response_text = re.sub(r'^```json?\n?', '', response_text)
                response_text = re.sub(r'\n?```$', '', response_text)
            
            dados = json.loads(response_text)
            itens = dados.get('itens', [])
            
            logger.info(f"✅ [DAMA] Extraídos {len(itens)} itens do edital")
            
            return itens
            
        except Exception as e:
            logger.error(f"❌ [DAMA] Erro ao extrair itens: {e}")
            # Retornar lista vazia em caso de erro
            return []
        finally:
            # Limpar arquivo temporário
            try:
                os.unlink(pdf_path)
            except Exception:
                pass
    
    async def extrair_itens_filtrado(self, pdf_content: bytes, palavra_chave: str) -> List[Dict]:
        """
        v73.0: Extrai itens do edital FILTRADOS por palavra-chave.
        
        Usa Gemini para identificar apenas itens que contenham o termo especificado.
        Ex: "Canabidiol" retorna apenas itens relacionados a canabidiol.
        
        Retorna lista de itens com campos para o usuário configurar:
        - numero: número do item no edital
        - descricao: descrição técnica
        - quantidade: quantidade licitada
        - unidade: unidade de medida
        """
        from emergentintegrations.llm.chat import LlmChat, UserMessage, FileContentWithMimeType
        
        # Salvar PDF temporariamente
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
            tmp.write(pdf_content)
            pdf_path = tmp.name
        
        # v73.1: Processar múltiplos termos separados por vírgula
        termos = [t.strip() for t in palavra_chave.split(',') if t.strip()]
        termos_str = ', '.join(termos)
        termos_lista = ' OU '.join([f'"{t}"' for t in termos])
        
        try:
            chat = LlmChat(
                api_key=self.api_key,
                session_id=f"dama-filtro-{datetime.now().timestamp()}",
                system_message=f"""Você é um especialista em análise de editais de licitação.
Sua tarefa é extrair APENAS os itens que contenham QUALQUER UMA das palavras-chave: {termos_lista}

SEMPRE responda em formato JSON válido com a seguinte estrutura:
{{
    "itens": [
        {{
            "numero": "1",
            "item": "1",
            "descricao": "Descrição técnica completa do item",
            "quantidade": 100,
            "unidade": "UN"
        }}
    ]
}}

REGRAS IMPORTANTES:
- Retorne itens que contenham QUALQUER UM dos termos: {termos_str}
- Se um item contiver "Canabidiol" OU "Insulina" OU qualquer outro termo da lista, inclua-o
- Se não encontrar nenhum item relacionado a nenhum dos termos, retorne {{"itens": []}}
- O campo "numero" e "item" devem ser o número do item no edital original
- Mantenha a descrição técnica completa
- Unidades comuns: UN, CX, FR, PCT, KG, L, ML, M, AMP, COMP, CP, FRA"""
            ).with_model("gemini", "gemini-2.5-flash")
            
            pdf_file = FileContentWithMimeType(
                file_path=pdf_path,
                mime_type="application/pdf"
            )
            
            response = await chat.send_message(UserMessage(
                text=f"Extraia TODOS os itens que contenham QUALQUER UM destes termos: {termos_lista}. Responda em JSON.",
                file_contents=[pdf_file]
            ))
            
            import json
            
            response_text = response.strip()
            if response_text.startswith('```'):
                response_text = re.sub(r'^```json?\n?', '', response_text)
                response_text = re.sub(r'\n?```$', '', response_text)
            
            dados = json.loads(response_text)
            itens = dados.get('itens', [])
            
            logger.info(f"✅ [DAMA v73.1] Extraídos {len(itens)} itens com filtro: {termos_str}")
            
            return itens
            
        except Exception as e:
            logger.error(f"❌ [DAMA v73.1] Erro ao extrair itens filtrados: {e}")
            return []
        finally:
            try:
                os.unlink(pdf_path)
            except Exception:
                pass
    
    async def _extrair_dados_pdf(self, pdf_content: bytes) -> DadosEdital:
        """
        Extrai dados do PDF usando Gemini 2.5 Flash.
        
        Usa FileContentWithMimeType para enviar o PDF diretamente ao modelo.
        """
        from emergentintegrations.llm.chat import LlmChat, UserMessage, FileContentWithMimeType
        
        # Salvar PDF temporariamente
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
            tmp.write(pdf_content)
            pdf_path = tmp.name
        
        try:
            chat = LlmChat(
                api_key=self.api_key,
                session_id=f"dama-ocr-{datetime.now().timestamp()}",
                system_message="""Você é um especialista em análise de editais de licitação pública no Brasil.
Sua tarefa é extrair informações estruturadas de editais em PDF.

SEMPRE responda em formato JSON válido com a seguinte estrutura:
{
    "orgao": "Nome completo do órgão licitante",
    "cnpj_orgao": "XX.XXX.XXX/XXXX-XX",
    "endereco_orgao": "Endereço completo",
    "uf_destino": "XX (sigla do estado)",
    "numero_processo": "Número do processo/licitação",
    "modalidade": "Pregão Eletrônico/Presencial/Concorrência/etc",
    "objeto": "Descrição resumida do objeto da licitação",
    "itens": [
        {
            "numero": "1",
            "descricao": "Descrição do item",
            "unidade": "UN/CX/FR/etc",
            "quantidade": 100,
            "valor_unitario_ref": 10.50,
            "valor_total_ref": 1050.00
        }
    ],
    "data_abertura": "DD/MM/AAAA HH:MM",
    "valor_referencia_total": 0.00,
    "observacoes": "Informações adicionais relevantes"
}

Se algum campo não for encontrado, use "N/A" para texto ou 0 para números.
Extraia TODOS os itens do edital, não apenas os primeiros."""
            ).with_model("gemini", "gemini-2.5-flash")
            
            # Criar objeto de arquivo
            pdf_file = FileContentWithMimeType(
                file_path=pdf_path,
                mime_type="application/pdf"
            )
            
            # Enviar mensagem com PDF
            response = await chat.send_message(UserMessage(
                text="Analise este edital de licitação e extraia todas as informações no formato JSON especificado.",
                file_contents=[pdf_file]
            ))
            
            # Parsear resposta JSON
            import json
            
            # Limpar resposta (remover markdown se houver)
            response_text = response.strip()
            if response_text.startswith('```'):
                response_text = re.sub(r'^```json?\n?', '', response_text)
                response_text = re.sub(r'\n?```$', '', response_text)
            
            dados = json.loads(response_text)
            
            # Converter para DadosEdital
            return DadosEdital(
                orgao=dados.get('orgao', 'Órgão não identificado'),
                cnpj_orgao=dados.get('cnpj_orgao', 'N/A'),
                endereco_orgao=dados.get('endereco_orgao', 'N/A'),
                uf_destino=dados.get('uf_destino', 'SP'),
                numero_processo=dados.get('numero_processo', 'N/A'),
                modalidade=dados.get('modalidade', 'Pregão Eletrônico'),
                objeto=dados.get('objeto', 'N/A'),
                itens=dados.get('itens', []),
                data_abertura=dados.get('data_abertura', 'N/A'),
                valor_referencia_total=float(dados.get('valor_referencia_total', 0)),
                observacoes=dados.get('observacoes', '')
            )
            
        finally:
            # Limpar arquivo temporário
            if os.path.exists(pdf_path):
                os.remove(pdf_path)
    
    def _calcular_tributacao(
        self,
        uf_origem: str,
        uf_destino: str,
        valor_total: float,
        custo_unitario: float
    ) -> Dict:
        """
        Calcula a tributação interestadual e margem sugerida.
        
        Regra de ICMS interestadual:
        - Sul/Sudeste para Sul/Sudeste: 12%
        - Sul/Sudeste para Norte/Nordeste/Centro-Oeste: 7%
        - Outras origens para Sul/Sudeste: 12%
        """
        # Buscar alíquota
        tabela_origem = ICMS_INTERESTADUAL.get(uf_origem, ICMS_INTERESTADUAL['DEFAULT'])
        icms_percent = tabela_origem.get(uf_destino, tabela_origem['DEFAULT'])
        
        # Calcular valores
        icms_valor = valor_total * (icms_percent / 100)
        valor_liquido = valor_total - icms_valor
        
        # Calcular margem se custo informado
        margem_liquida = 0
        margem_percent = 0
        if custo_unitario > 0 and valor_total > 0:
            margem_liquida = valor_liquido - custo_unitario
            margem_percent = (margem_liquida / custo_unitario) * 100
        
        return {
            'uf_origem': uf_origem,
            'uf_destino': uf_destino,
            'icms_percent': icms_percent,
            'icms_valor': icms_valor,
            'valor_bruto': valor_total,
            'valor_liquido': valor_liquido,
            'custo': custo_unitario,
            'margem_liquida': margem_liquida,
            'margem_percent': margem_percent
        }
    
    async def _gerar_texto_proposta(
        self,
        dados: DadosEdital,
        empresa: EmpresaGSM,
        calculo: Dict,
        custo_unitario: float
    ) -> str:
        """
        Gera o texto da proposta comercial usando Gemini.
        """
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        
        # Formatar itens para o prompt
        itens_texto = "\n".join([
            f"- Item {it.get('numero', i+1)}: {it.get('descricao', 'Item')} | "
            f"Qtd: {it.get('quantidade', 0)} {it.get('unidade', 'UN')} | "
            f"Valor Ref: R$ {it.get('valor_total_ref', 0):,.2f}"
            for i, it in enumerate(dados.itens)
        ])
        
        prompt = f"""Gere uma proposta comercial profissional para licitação pública.

DADOS DO EDITAL:
- Órgão: {dados.orgao}
- Processo: {dados.numero_processo}
- Modalidade: {dados.modalidade}
- Objeto: {dados.objeto}
- Data de Abertura: {dados.data_abertura}
- UF Destino: {dados.uf_destino}

ITENS:
{itens_texto}

DADOS DA EMPRESA PROPONENTE:
- Razão Social: {empresa.nome}
- CNPJ: {empresa.cnpj}
- Endereço: {empresa.endereco}
- Inscrição Estadual: {empresa.inscricao_estadual}
- Telefone: {empresa.telefone}
- E-mail: {empresa.email}

INFORMAÇÕES TRIBUTÁRIAS:
- UF Origem: {calculo['uf_origem']}
- ICMS Aplicável: {calculo['icms_percent']}%
- Valor de Referência: R$ {calculo['valor_bruto']:,.2f}

Gere uma proposta comercial completa incluindo:
1. Identificação do processo
2. Apresentação da empresa
3. Tabela de itens com preços (use valores de referência como base, podendo aplicar desconto de 1-5%)
4. Condições de pagamento (conforme edital)
5. Prazo de entrega
6. Validade da proposta (60 dias)
7. Declaração de conhecimento e concordância com o edital

FORMATO: Texto corrido em linguagem formal de licitação, pronto para inserir em documento Word.
NÃO inclua cabeçalho ou rodapé (serão mantidos do template do cliente).
"""
        
        chat = LlmChat(
            api_key=self.api_key,
            session_id=f"dama-proposta-{datetime.now().timestamp()}",
            system_message="Você é um especialista em elaboração de propostas comerciais para licitações públicas no Brasil. Gere textos profissionais, formais e completos."
        ).with_model("gemini", "gemini-2.5-flash")
        
        response = await chat.send_message(UserMessage(text=prompt))
        
        return response
    
    async def _gerar_texto_proposta_simples(
        self,
        dados: DadosEdital,
        empresa: EmpresaGSM
    ) -> str:
        """
        v73.1: Gera o texto da proposta comercial SEM cálculo tributário.
        """
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        
        # Formatar itens para o prompt
        itens_texto = "\n".join([
            f"- Item {it.get('numero', i+1)}: {it.get('descricao', 'Item')} | "
            f"Qtd: {it.get('quantidade', 0)} {it.get('unidade', 'UN')} | "
            f"Preço Unit: R$ {it.get('preco_unitario', 0):,.2f}"
            for i, it in enumerate(dados.itens)
        ])
        
        prompt = f"""Gere uma proposta comercial profissional para licitação pública.

DADOS DO EDITAL:
- Órgão: {dados.orgao}
- Processo: {dados.numero_processo}
- Modalidade: {dados.modalidade}
- Objeto: {dados.objeto}
- Data de Abertura: {dados.data_abertura}
- UF Destino: {dados.uf_destino}

ITENS COTADOS:
{itens_texto}

DADOS DA EMPRESA PROPONENTE:
- Razão Social: {empresa.nome}
- CNPJ: {empresa.cnpj}
- Endereço: {empresa.endereco}
- Inscrição Estadual: {empresa.inscricao_estadual}
- Telefone: {empresa.telefone}
- E-mail: {empresa.email}

Gere uma proposta comercial completa incluindo:
1. Identificação do processo licitatório
2. Apresentação da empresa proponente
3. Tabela de itens com os preços informados
4. Condições de pagamento (conforme edital)
5. Prazo de entrega
6. Validade da proposta (60 dias)
7. Declaração de conhecimento e concordância com o edital

FORMATO: Texto corrido em linguagem formal de licitação, pronto para inserir em documento Word.
NÃO inclua cabeçalho ou rodapé (serão mantidos do template do cliente).
"""
        
        chat = LlmChat(
            api_key=self.api_key,
            session_id=f"dama-proposta-{datetime.now().timestamp()}",
            system_message="Você é um especialista em elaboração de propostas comerciais para licitações públicas no Brasil. Gere textos profissionais, formais e completos."
        ).with_model("gemini", "gemini-2.5-flash")
        
        response = await chat.send_message(UserMessage(text=prompt))
        
        return response
    
    def _injetar_template(
        self,
        template_content: bytes,
        texto_proposta: str,
        dados: DadosEdital,
        empresa: EmpresaGSM
    ) -> bytes:
        """
        Injeta o texto da proposta no template Word.
        
        Localiza a tag {{TEXTO_DAMA}} e substitui pelo texto gerado.
        Preserva cabeçalho, rodapé e formatação do template original.
        """
        # Carregar template
        doc = Document(io.BytesIO(template_content))
        
        # Substituir tag {{TEXTO_DAMA}} em parágrafos
        tag_encontrada = False
        for para in doc.paragraphs:
            if '{{TEXTO_DAMA}}' in para.text:
                tag_encontrada = True
                # Preservar formatação do parágrafo
                para.clear()
                # Adicionar texto da proposta
                for linha in texto_proposta.split('\n'):
                    run = para.add_run(linha + '\n')
                    run.font.size = Pt(11)
                    run.font.name = 'Arial'
        
        # Se não encontrou em parágrafos, buscar em tabelas
        if not tag_encontrada:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if '{{TEXTO_DAMA}}' in cell.text:
                            tag_encontrada = True
                            cell.text = texto_proposta
        
        # Se ainda não encontrou, adicionar ao final
        if not tag_encontrada:
            logger.warning("⚠️ [DAMA] Tag {{TEXTO_DAMA}} não encontrada, adicionando ao final")
            doc.add_paragraph(texto_proposta)
        
        # Substituir outras tags opcionais
        for para in doc.paragraphs:
            if '{{EMPRESA}}' in para.text:
                para.text = para.text.replace('{{EMPRESA}}', empresa.nome)
            if '{{CNPJ}}' in para.text:
                para.text = para.text.replace('{{CNPJ}}', empresa.cnpj)
            if '{{PROCESSO}}' in para.text:
                para.text = para.text.replace('{{PROCESSO}}', dados.numero_processo)
            if '{{ORGAO}}' in para.text:
                para.text = para.text.replace('{{ORGAO}}', dados.orgao)
            if '{{DATA}}' in para.text:
                para.text = para.text.replace('{{DATA}}', datetime.now().strftime('%d/%m/%Y'))
        
        # Salvar em bytes
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()
    
    def _gerar_declaracoes(self, empresa: EmpresaGSM, dados: DadosEdital, timbrado_content: bytes = None) -> Dict[str, bytes]:
        """
        v65.1: Gera declarações obrigatórias USANDO O TIMBRADO da empresa.
        
        Se timbrado_content for fornecido, usa como base para as declarações.
        Caso contrário, gera documentos simples.
        
        Declarações:
        1. Declaração de Inexistência de Fato Impeditivo
        2. Declaração de Não Emprego de Menor
        3. Declaração de Elaboração Independente de Proposta
        """
        declaracoes = {}
        
        # Textos das declarações
        texto_declaracao_1 = f"""DECLARAÇÃO DE INEXISTÊNCIA DE FATO IMPEDITIVO

Declaramos, sob as penas da lei, que a empresa {empresa.nome}, inscrita no CNPJ sob o nº {empresa.cnpj}, com sede em {empresa.endereco}, não se encontra impedida de participar de licitações ou contratar com a Administração Pública.

Declaramos, ainda, que não existe qualquer fato impeditivo à nossa habilitação e que nos comprometemos a comunicar a ocorrência de qualquer fato superveniente.

Local e Data: __________________, {datetime.now().strftime('%d de %B de %Y')}


_________________________________________________
{empresa.nome}
CNPJ: {empresa.cnpj}"""

        texto_declaracao_2 = f"""DECLARAÇÃO DE NÃO EMPREGO DE MENOR

Ref.: {dados.modalidade} - Processo nº {dados.numero_processo}

A empresa {empresa.nome}, inscrita no CNPJ sob o nº {empresa.cnpj}, por intermédio de seu representante legal, DECLARA, para fins do disposto no inciso V do art. 27 da Lei Federal nº 8.666/93, acrescido pela Lei nº 9.854/99, que não emprega menor de dezoito anos em trabalho noturno, perigoso ou insalubre, e não emprega menor de dezesseis anos.

Ressalva: emprega menor, a partir de quatorze anos, na condição de aprendiz ( ).

Local e Data: __________________, {datetime.now().strftime('%d de %B de %Y')}


_________________________________________________
Representante Legal
{empresa.nome}"""

        texto_declaracao_3 = f"""DECLARAÇÃO DE ELABORAÇÃO INDEPENDENTE DE PROPOSTA

Ref.: {dados.modalidade} - Processo nº {dados.numero_processo}
Órgão: {dados.orgao}

A empresa {empresa.nome}, inscrita no CNPJ sob o nº {empresa.cnpj}, por intermédio de seu representante legal, DECLARA, sob as penas da lei, em especial o art. 299 do Código Penal Brasileiro, que:

a) A proposta apresentada para participar desta licitação foi elaborada de maneira independente, e o conteúdo da proposta não foi, no todo ou em parte, direta ou indiretamente, informado, discutido ou recebido de qualquer outro participante potencial ou de fato desta licitação;

b) A intenção de apresentar a proposta elaborada para participar desta licitação não foi informada, discutida ou recebida de qualquer outro participante potencial ou de fato desta licitação;

c) Não tentou, por qualquer meio ou por qualquer pessoa, influir na decisão de qualquer outro participante potencial ou de fato desta licitação quanto a participar ou não da referida licitação.

Local e Data: __________________, {datetime.now().strftime('%d de %B de %Y')}


_________________________________________________
Representante Legal
{empresa.nome}
CNPJ: {empresa.cnpj}"""

        declaracoes_config = [
            ('declaracao_inexistencia_fato_impeditivo.docx', texto_declaracao_1),
            ('declaracao_nao_emprego_menor.docx', texto_declaracao_2),
            ('declaracao_elaboracao_independente.docx', texto_declaracao_3)
        ]
        
        for filename, texto in declaracoes_config:
            try:
                if timbrado_content:
                    # v65.1: Usar timbrado como base
                    doc = Document(io.BytesIO(timbrado_content))
                    
                    # Procurar e substituir a tag {{TEXTO_DAMA}}
                    tag_encontrada = False
                    for para in doc.paragraphs:
                        if '{{TEXTO_DAMA}}' in para.text:
                            para.text = para.text.replace('{{TEXTO_DAMA}}', texto)
                            tag_encontrada = True
                            break
                    
                    # Se não encontrou a tag, adicionar ao final
                    if not tag_encontrada:
                        doc.add_paragraph(texto)
                    
                    # Substituir outras tags opcionais
                    for para in doc.paragraphs:
                        if '{{EMPRESA}}' in para.text:
                            para.text = para.text.replace('{{EMPRESA}}', empresa.nome)
                        if '{{CNPJ}}' in para.text:
                            para.text = para.text.replace('{{CNPJ}}', empresa.cnpj)
                        if '{{PROCESSO}}' in para.text:
                            para.text = para.text.replace('{{PROCESSO}}', dados.numero_processo)
                        if '{{ORGAO}}' in para.text:
                            para.text = para.text.replace('{{ORGAO}}', dados.orgao)
                        if '{{DATA}}' in para.text:
                            para.text = para.text.replace('{{DATA}}', datetime.now().strftime('%d/%m/%Y'))
                else:
                    # Sem timbrado: gerar documento simples
                    doc = Document()
                    for linha in texto.split('\n'):
                        doc.add_paragraph(linha)
                
                output = io.BytesIO()
                doc.save(output)
                declaracoes[filename] = output.getvalue()
                logger.info(f"✅ [DAMA v65.1] Declaração gerada: {filename} (com timbrado: {bool(timbrado_content)})")
                
            except Exception as e:
                logger.error(f"❌ [DAMA] Erro ao gerar {filename}: {e}")
                # Fallback: documento simples
                doc = Document()
                for linha in texto.split('\n'):
                    doc.add_paragraph(linha)
                output = io.BytesIO()
                doc.save(output)
                declaracoes[filename] = output.getvalue()
        
        return declaracoes
    
    def _montar_zip(
        self,
        proposta_docx: bytes,
        declaracoes: Dict[str, bytes],
        dados: DadosEdital
    ) -> bytes:
        """
        Monta o ZIP final com todos os documentos.
        
        Estrutura:
        - kit_licitacao_PROCESSO.zip
          ├── PROPOSTA_COMERCIAL.docx
          ├── DECLARACOES/
          │   ├── declaracao_inexistencia_fato_impeditivo.docx
          │   ├── declaracao_nao_emprego_menor.docx
          │   └── declaracao_elaboracao_independente.docx
        """
        zip_buffer = io.BytesIO()
        
        # Limpar número do processo para nome do arquivo
        processo_limpo = re.sub(r'[^\w\-]', '_', dados.numero_processo)
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            # Adicionar proposta
            zf.writestr(f'PROPOSTA_COMERCIAL_{processo_limpo}.docx', proposta_docx)
            
            # Adicionar declarações
            for nome_arquivo, conteudo in declaracoes.items():
                zf.writestr(f'DECLARACOES/{nome_arquivo}', conteudo)
        
        return zip_buffer.getvalue()


# Singleton
_dama_service = None

def get_dama_service() -> DAMAService:
    global _dama_service
    if _dama_service is None:
        _dama_service = DAMAService()
    return _dama_service
